# 接口1: 文档上传与智能分割
# POST /api/data/upload_split

from flask import Flask, request, jsonify
from splitter import DocumentSplitter
import io
import os
import re
import sys
import time
import uuid
import zipfile

# 添加shared模块路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from shared.database.mysql_client import MySQLClient
from shared.database.models import Block
from shared.utils.file_store import FileStore
from protocol_extractor import enrich_protocol_metadata

app = Flask(__name__)
splitter = DocumentSplitter()

# 初始化共享模块
mysql_client = MySQLClient()
file_store = FileStore()

UPLOAD_DIR = "/tmp/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

SUPPORTED_EXTS = [
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
    ".txt",
    ".md",
    ".py",
    ".java",
    ".js",
    ".cpp",
    ".json",
    ".xml",
    ".yaml",
    ".yml",
]
TEXT_BASED_EXTS = {".txt", ".md", ".py", ".java", ".js", ".cpp", ".json", ".xml", ".yaml", ".yml"}
DEFAULT_MAX_FILE_SIZE_MB = 50
MIN_READABLE_CHAR_COUNT = 20
ZIP_SIGNATURES = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")
OLE_SIGNATURE = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"


def _normalize_cell(value) -> str:
    """标准化单元格文本"""
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _normalize_table_rows(table: list) -> list:
    """标准化表格行，移除空行并补齐列数"""
    rows = []
    max_cols = 0
    for row in table or []:
        cleaned = [_normalize_cell(cell) for cell in (row or [])]
        if any(cleaned):
            rows.append(cleaned)
            max_cols = max(max_cols, len(cleaned))

    if max_cols == 0:
        return []

    normalized = []
    for row in rows:
        if len(row) < max_cols:
            normalized.append(row + [""] * (max_cols - len(row)))
        else:
            normalized.append(row[:max_cols])
    return normalized


def _row_signature(row: list) -> str:
    return "|".join([_normalize_cell(cell).lower() for cell in row if _normalize_cell(cell)])


def _row_blank_ratio(row: list) -> float:
    if not row:
        return 1.0
    blank = sum(1 for cell in row if not _normalize_cell(cell))
    return blank / max(1, len(row))


def _cell_shape(cell: str) -> str:
    text = _normalize_cell(cell)
    if not text:
        return "e"
    has_digit = bool(re.search(r"\d", text))
    has_alpha = bool(re.search(r"[A-Za-z\u4e00-\u9fff]", text))
    if has_alpha and has_digit:
        return "m"
    if has_digit:
        return "n"
    if has_alpha:
        return "a"
    return "o"


def _row_shape_similarity(left: list, right: list) -> float:
    size = max(len(left), len(right), 1)
    score = 0
    for idx in range(size):
        left_shape = _cell_shape(left[idx] if idx < len(left) else "")
        right_shape = _cell_shape(right[idx] if idx < len(right) else "")
        if left_shape == right_shape:
            score += 1
    return score / size


def _looks_like_header_row(row: list) -> bool:
    non_empty = [_normalize_cell(cell) for cell in row if _normalize_cell(cell)]
    if len(non_empty) < 2:
        return False

    joined = " ".join(non_empty).lower()
    hints = ("field", "字段", "name", "名称", "bit", "位", "desc", "描述", "length", "len", "type", "unit")
    if any(keyword in joined for keyword in hints):
        return True

    alpha_cells = sum(1 for cell in non_empty if re.search(r"[A-Za-z\u4e00-\u9fff]", cell))
    digit_cells = sum(1 for cell in non_empty if re.search(r"\d", cell))
    return alpha_cells >= max(1, len(non_empty) - 1) and digit_cells <= len(non_empty) // 2


def _parse_text_tables(text: str, page_num: int) -> list:
    """从文本中解析伪表格（Word Map格式等）"""
    tables = []

    # Word Map格式检测
    if "WORD NUMBER:" in text or "WORD MAP" in text:
        # 提取Word Map块
        word_maps = re.split(r'WORD NUMBER:', text)
        for wm in word_maps[1:]:  # 跳过第一个空块
            lines = wm.strip().split('\n')
            if len(lines) < 3:
                continue

            # 提取Word Number
            word_num_match = re.match(r'\s*(J[\d.]+[IE]?)', lines[0])
            word_num = word_num_match.group(1) if word_num_match else "Unknown"

            # 提取Word Title
            title = ""
            for i, line in enumerate(lines):
                if "WORD TITLE:" in line:
                    title_match = re.search(r'WORD TITLE:\s*(.+)', line)
                    if title_match:
                        title = title_match.group(1).strip()
                    break

            # 查找表格行（包含冒号分隔符的行）
            table_rows = []
            for line in lines:
                # 检测表格行模式：多个冒号分隔的字段
                if line.count(':') >= 2 and not line.strip().startswith('---'):
                    cells = [c.strip() for c in line.split(':') if c.strip()]
                    if cells:
                        table_rows.append(cells)

            if table_rows:
                tables.append({
                    "page_num": page_num,
                    "content": f"Word Map: {word_num}\nTitle: {title}\n" + format_table_to_text(table_rows),
                    "type": "table",
                    "metadata": {
                        "word_number": word_num,
                        "word_title": title,
                        "row_count": len(table_rows),
                        "col_count": max(len(r) for r in table_rows) if table_rows else 0,
                        "table_type": "word_map",
                    },
                })

    # Bit字段表格检测（格式: Field | Bit | Meaning）
    bit_pattern = re.compile(r'^\s*([A-Z_][A-Z0-9_]*)\s*[\|\t]\s*(\d+)\s*[-–~]\s*(\d+)\s*[\|\t]\s*(.+)$', re.MULTILINE)
    bit_matches = bit_pattern.findall(text)
    if bit_matches:
        rows = []
        for match in bit_matches:
            field_name, bit_start, bit_end, meaning = match
            rows.append([field_name, f"{bit_start}-{bit_end}", meaning.strip()])
        if rows:
            tables.append({
                "page_num": page_num,
                "content": "Bit Field Table\n" + format_table_to_text(rows),
                "type": "table",
                "metadata": {
                    "row_count": len(rows),
                    "col_count": 3,
                    "table_type": "bit_field",
                },
            })

    # J消息格式检测（J2.0, J3.0等）
    j_msg_pattern = re.compile(r'(J[\d.]+[IE]?)\s*[\|\t:]\s*(.+?)(?:\s*[\|\t]\s*(.+))?$', re.MULTILINE)
    j_matches = j_msg_pattern.findall(text)
    if len(j_matches) >= 3:  # 至少3个才认为是表格
        rows = []
        for match in j_matches[:20]:  # 限制数量
            msg_code = match[0]
            desc = match[1].strip() if len(match) > 1 else ""
            extra = match[2].strip() if len(match) > 2 else ""
            rows.append([msg_code, desc, extra] if extra else [msg_code, desc])
        if rows:
            tables.append({
                "page_num": page_num,
                "content": "J-Message Format\n" + format_table_to_text(rows),
                "type": "table",
                "metadata": {
                    "row_count": len(rows),
                    "col_count": max(len(r) for r in rows),
                    "table_type": "j_message",
                },
            })

    return tables


def _extract_page_tables(page, page_num: int) -> list:
    """提取页内表格并保留跨页拼接所需信息"""
    table_blocks = []
    page_height = float(page.height or 1)

    # 方法1：pdfplumber标准表格提取
    for table_idx, table_obj in enumerate(page.find_tables() or []):
        rows = _normalize_table_rows(table_obj.extract() if hasattr(table_obj, "extract") else [])
        if not rows:
            continue
        bbox = getattr(table_obj, "bbox", None) or (0.0, 0.0, 0.0, page_height)
        top = float(bbox[1])
        bottom = float(bbox[3])
        table_blocks.append(
            {
                "page_num": page_num,
                "content": format_table_to_text(rows),
                "type": "table",
                "metadata": {
                    "row_count": len(rows),
                    "col_count": max(len(r) for r in rows),
                    "near_top": top <= page_height * 0.22,
                    "near_bottom": bottom >= page_height * 0.78,
                    "source_pages": [page_num],
                    "table_index": table_idx,
                    "last_page_num": page_num,
                },
                "_rows": rows,
                "_order": (page_num, 0, table_idx),
            }
        )

    # 方法2：从文本中解析伪表格（Word Map等格式）
    text = page.extract_text() or ""
    text_tables = _parse_text_tables(text, page_num)
    for idx, tt in enumerate(text_tables):
        table_blocks.append({
            "page_num": page_num,
            "content": tt["content"],
            "type": "table",
            "metadata": tt.get("metadata", {}),
            "_order": (page_num, 0, len(table_blocks) + idx),
        })

    return table_blocks


def _should_merge_cross_page_table(previous: dict, current: dict) -> tuple:
    """判断两个表格是否属于跨页续接"""
    previous_meta = previous.get("metadata", {})
    current_meta = current.get("metadata", {})

    if current["page_num"] != previous_meta.get("last_page_num", previous["page_num"]) + 1:
        return False, False
    if not previous_meta.get("near_bottom") or not current_meta.get("near_top"):
        return False, False

    previous_cols = int(previous_meta.get("col_count", 0))
    current_cols = int(current_meta.get("col_count", 0))
    if previous_cols <= 0 or current_cols <= 0 or abs(previous_cols - current_cols) > 1:
        return False, False

    previous_rows = previous.get("_rows", [])
    current_rows = current.get("_rows", [])
    if not previous_rows or not current_rows:
        return False, False

    repeated_header = _row_signature(current_rows[0]) == _row_signature(previous_rows[0]) and bool(_row_signature(current_rows[0]))
    if repeated_header:
        return True, True

    current_starts_with_header = _looks_like_header_row(current_rows[0])
    if current_starts_with_header:
        return False, False

    similarity = _row_shape_similarity(previous_rows[-1], current_rows[0])
    sparse_tail = _row_blank_ratio(previous_rows[-1]) >= 0.4
    if previous_cols == current_cols and (similarity >= 0.75 or (sparse_tail and similarity >= 0.6)):
        return True, False
    return False, False


def _merge_cross_page_tables(table_blocks: list) -> list:
    """合并跨页续接表格，避免拆断"""
    if not table_blocks:
        return []

    merged = []
    for current in table_blocks:
        if merged:
            can_merge, drop_header = _should_merge_cross_page_table(merged[-1], current)
            if can_merge:
                append_rows = current.get("_rows", [])
                if drop_header and len(append_rows) > 1:
                    append_rows = append_rows[1:]
                elif drop_header and len(append_rows) <= 1:
                    append_rows = []

                if append_rows:
                    merged[-1]["_rows"].extend(append_rows)

                source_pages = merged[-1]["metadata"].setdefault("source_pages", [merged[-1]["page_num"]])
                for page_num in current["metadata"].get("source_pages", [current["page_num"]]):
                    if page_num not in source_pages:
                        source_pages.append(page_num)

                merged[-1]["content"] = format_table_to_text(merged[-1]["_rows"])
                merged[-1]["metadata"]["row_count"] = len(merged[-1]["_rows"])
                merged[-1]["metadata"]["col_count"] = max(
                    int(merged[-1]["metadata"].get("col_count", 0)),
                    int(current["metadata"].get("col_count", 0)),
                )
                merged[-1]["metadata"]["cross_page_merged"] = len(source_pages) > 1
                merged[-1]["metadata"]["last_page_num"] = max(source_pages)
                merged[-1]["metadata"]["near_bottom"] = bool(current["metadata"].get("near_bottom"))
                continue
        merged.append(current)

    for table_block in merged:
        table_block.pop("_rows", None)
        table_block["metadata"].pop("last_page_num", None)
    return merged


def process_pdf_with_pages(file_path: str, enable_llm_postprocess: bool = True) -> list:
    """
    处理PDF文件，按页解析并检测表格
    返回: [{"page_num": int, "content": str, "type": "text|table"}, ...]
    """
    import pdfplumber

    text_blocks = []
    table_blocks = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            table_blocks.extend(_extract_page_tables(page, page_num))

            # 提取文本内容（排除表格区域）
            text = page.extract_text()
            if text and text.strip():
                text_blocks.append(
                    {
                        "page_num": page_num,
                        "content": text.strip(),
                        "type": "text",
                        "metadata": {},
                        "_order": (page_num, 1, 0),
                    }
                )

    merged_table_blocks = _merge_cross_page_tables(table_blocks)
    blocks = merged_table_blocks + text_blocks
    blocks.sort(key=lambda item: item.get("_order", (item["page_num"], 1, 0)))

    enriched_blocks = []
    for block in blocks:
        normalized_block = {
            "page_num": block["page_num"],
            "content": block["content"],
            "type": block["type"],
            "metadata": block.get("metadata", {}),
        }
        enriched_blocks.append(enrich_protocol_metadata(normalized_block, enable_llm_postprocess))

    return enriched_blocks


def format_table_to_text(table: list) -> str:
    """将表格数据转换为格式化文本"""
    if not table:
        return ""

    lines = []
    for row in table:
        # 处理空值，转换为空字符串
        formatted_row = [str(cell) if cell is not None else "" for cell in row]
        lines.append(" | ".join(formatted_row))

    return "\n".join(lines)


def process_docx_with_pages(file_path: str, enable_llm_postprocess: bool = True) -> list:
    """
    处理DOCX文件，检测表格并分块
 返回: [{"page_num": int, "content": str, "type": "text|table"}, ...]
    """
    from docx import Document

    blocks = []
    doc = Document(file_path)

    # DOCX没有明确的页码概念，按段落和表格顺序处理
    current_text = []
    block_order = 0

    # 遍历文档元素（段落和表格）
    for element in doc.element.body:
        # 处理表格
        if element.tag.endswith('tbl'):
            # 先保存之前的文本
            if current_text:
                text_content = "\n".join(current_text).strip()
                if text_content:
                    block_order += 1
                    blocks.append({
                        "page_num": block_order,  # DOCX用block_order模拟页码
                        "content": text_content,
                        "type": "text",
                        "metadata": {}
                    })
                current_text = []

            # 处理表格
            table = None
            for tbl in doc.tables:
                if tbl._tbl == element:
                    table = tbl
                    break

            if table:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                table_text = format_table_to_text(table_data)
                if table_text.strip():
                    block_order += 1
                    blocks.append({
                        "page_num": block_order,
                        "content": table_text,
                        "type": "table",
                        "metadata": {"row_count": len(table_data), "col_count": len(table_data[0]) if table_data else 0}
                    })

        # 处理段落
        elif element.tag.endswith('p'):
            for para in doc.paragraphs:
                if para._p == element:
                    if para.text.strip():
                        current_text.append(para.text.strip())
                    break

    # 保存剩余的文本
    if current_text:
        text_content = "\n".join(current_text).strip()
        if text_content:
            block_order += 1
            blocks.append({
                "page_num": block_order,
                "content": text_content,
                "type": "text",
                "metadata": {}
            })

    return [enrich_protocol_metadata(block, enable_llm_postprocess) for block in blocks]


def process_excel_with_sheets(file_path: str, enable_llm_postprocess: bool = True) -> list:
    """
    处理Excel文件，按sheet解析为可分块文本/表格块
    返回: [{"page_num": int, "content": str, "type": "text|table"}, ...]
    """
    ext = os.path.splitext(file_path)[1].lower()
    sheets = []

    if ext == ".xlsx":
        from openpyxl import load_workbook

        workbook = load_workbook(file_path, data_only=True, read_only=True)
        try:
            for index, sheet in enumerate(workbook.worksheets, start=1):
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    cleaned = [_normalize_cell(cell) for cell in row]
                    if any(cleaned):
                        rows.append(cleaned)
                sheets.append((index, sheet.title, rows))
        finally:
            workbook.close()
    elif ext == ".xls":
        try:
            import xlrd
        except ImportError as exc:
            raise ValueError("未安装xlrd，无法解析.xls文件") from exc

        workbook = xlrd.open_workbook(file_path)
        for index in range(workbook.nsheets):
            sheet = workbook.sheet_by_index(index)
            rows = []
            for row_idx in range(sheet.nrows):
                cleaned = [_normalize_cell(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
                if any(cleaned):
                    rows.append(cleaned)
            sheets.append((index + 1, sheet.name, rows))
    else:
        raise ValueError(f"不支持的Excel文件类型: {ext}")

    blocks = []
    for page_num, sheet_name, rows in sheets:
        normalized_rows = _normalize_table_rows(rows)
        if not normalized_rows:
            continue

        col_count = max(len(row) for row in normalized_rows)
        if col_count <= 1:
            content = "\n".join(row[0] for row in normalized_rows if row and row[0]).strip()
            block_type = "text"
        else:
            content = format_table_to_text(normalized_rows)
            block_type = "table"

        if not content:
            continue

        block = {
            "page_num": page_num,
            "content": content,
            "type": block_type,
            "metadata": {
                "sheet_name": sheet_name,
                "sheet_index": page_num,
                "row_count": len(normalized_rows),
                "col_count": col_count,
            },
        }
        blocks.append(enrich_protocol_metadata(block, enable_llm_postprocess))

    return blocks


def split_large_content(
    content: str,
    page_num: int,
    block_type: str,
    file_ext: str,
    base_metadata: dict | None = None,
) -> list:
    """
    对大块内容进行分割，保持页码和类型信息
    返回: [{"page_num": int, "content": str, "type": str}, ...]
    """
    if file_ext == "pdf":
        splitter_type = "pdf"
    elif file_ext in {"docx", "xlsx", "xls"}:
        splitter_type = "docx"
    else:
        splitter_type = "txt"
    splitter_instance = splitter.get_splitter(doc_type=splitter_type)

    if hasattr(splitter_instance, "split_text"):
        chunks = splitter_instance.split_text(content)
    else:
        docs = splitter_instance.create_documents([content])
        chunks = [d.page_content for d in docs]

    result = []
    for idx, chunk in enumerate(chunks):
        if chunk.strip():
            metadata = dict(base_metadata or {})
            metadata.update({"chunk_index": idx, "total_chunks": len(chunks)})
            result.append({
                "page_num": page_num,
                "content": chunk.strip(),
                "type": block_type,
                "metadata": metadata
            })

    return result


def process_file(
    file_path: str,
    project_id: str,
    file_name: str,
    enable_llm_postprocess: bool = True,
) -> tuple:
    """
    处理文件，返回(blocks, total_pages)
    blocks: 用于响应的块列表
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        raw_blocks = process_pdf_with_pages(file_path, enable_llm_postprocess=enable_llm_postprocess)
    elif ext == ".docx":
        raw_blocks = process_docx_with_pages(file_path, enable_llm_postprocess=enable_llm_postprocess)
    elif ext in {".xlsx", ".xls"}:
        raw_blocks = process_excel_with_sheets(file_path, enable_llm_postprocess=enable_llm_postprocess)
    else:
        # 其他文件类型使用原有逻辑
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        documents = splitter.split_to_documents(content, file_path=file_path)
        raw_blocks = []
        for idx, doc in enumerate(documents):
            metadata = doc.metadata if isinstance(doc.metadata, dict) else {}
            block = {
                "page_num": idx + 1,
                "content": doc.page_content,
                "type": "text",
                "metadata": dict(metadata),
            }
            raw_blocks.append(enrich_protocol_metadata(block, enable_llm_postprocess))

    # 对大块内容进行分割
    final_blocks = []
    for raw_block in raw_blocks:
        content = raw_block["content"]
        # 如果内容较长，进行分割
        if len(content) > 1000:
            split_blocks = split_large_content(
                content,
                raw_block["page_num"],
                raw_block["type"],
                ext.lstrip("."),
                base_metadata=raw_block.get("metadata", {}),
            )
            final_blocks.extend(split_blocks)
        else:
            final_blocks.append(raw_block)

    # 计算总页数
    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
    else:
        # 对于非PDF文件，使用最大页码值
        total_pages = max([b["page_num"] for b in final_blocks]) if final_blocks else 0

    return final_blocks, total_pages


def _file_size_bytes(uploaded_file) -> int:
    current_pos = uploaded_file.stream.tell()
    uploaded_file.stream.seek(0, os.SEEK_END)
    size = uploaded_file.stream.tell()
    uploaded_file.stream.seek(current_pos)
    return size


def _read_uploaded_bytes(uploaded_file) -> bytes:
    uploaded_file.stream.seek(0)
    file_bytes = uploaded_file.read()
    uploaded_file.stream.seek(0)
    return file_bytes


def _looks_like_text_content(file_bytes: bytes) -> bool:
    if not file_bytes or b"\x00" in file_bytes:
        return False

    sample = file_bytes[:4096]
    printable = sum(
        1
        for byte in sample
        if byte in {9, 10, 13} or 32 <= byte <= 126 or byte >= 128
    )
    if printable / max(len(sample), 1) < 0.85:
        return False

    try:
        return bool(sample.decode("utf-8", errors="ignore").strip())
    except Exception:
        return False


def _detect_file_authenticity(ext: str, file_bytes: bytes) -> tuple[bool, str, str]:
    if ext == ".pdf":
        if file_bytes.startswith(b"%PDF-"):
            return True, "pdf", "文件头匹配 PDF 签名"
        return False, "unknown", "文件扩展名为 .pdf，但文件头未匹配 PDF 签名"

    if ext in {".docx", ".xlsx"}:
        if not any(file_bytes.startswith(signature) for signature in ZIP_SIGNATURES):
            return False, "unknown", f"文件扩展名为 {ext}，但文件头未匹配 Office ZIP 签名"
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
                members = set(archive.namelist())
        except zipfile.BadZipFile:
            return False, "unknown", "文件扩展名指向 Office 文档，但压缩结构已损坏"

        if ext == ".docx" and "word/document.xml" in members:
            return True, "docx", "文件扩展名与 DOCX 包结构一致"
        if ext == ".xlsx" and "xl/workbook.xml" in members:
            return True, "xlsx", "文件扩展名与 XLSX 包结构一致"

        detected_type = "xlsx" if "xl/workbook.xml" in members else "docx" if "word/document.xml" in members else "unknown"
        return False, detected_type, f"文件扩展名为 {ext}，但包内结构与实际类型不一致"

    if ext == ".xls":
        if file_bytes.startswith(OLE_SIGNATURE):
            return True, "xls", "文件头匹配 XLS 复合文档签名"
        if any(file_bytes.startswith(signature) for signature in ZIP_SIGNATURES):
            return False, "xlsx", "文件扩展名为 .xls，但内容更像 XLSX 压缩包"
        return False, "unknown", "文件扩展名为 .xls，但文件头未匹配 XLS 签名"

    if ext in TEXT_BASED_EXTS:
        if _looks_like_text_content(file_bytes):
            return True, "text", "文本类文件可正常解码且未检测到二进制特征"
        return False, "binary", f"文件扩展名为 {ext}，但内容更像二进制文件"

    return False, "unknown", f"暂未支持识别扩展名 {ext} 的真实性"


def _save_validation_temp_file(file_name: str, file_bytes: bytes) -> str:
    file_path = os.path.join(UPLOAD_DIR, f"validate_{uuid.uuid4().hex}_{file_name}")
    with open(file_path, "wb") as temp_file:
        temp_file.write(file_bytes)
    return file_path


def _analyze_file_content(file_path: str, file_name: str) -> tuple[dict, list[str]]:
    issues = []
    raw_blocks, total_pages = process_file(
        file_path,
        project_id="validation_preview",
        file_name=file_name,
        enable_llm_postprocess=False,
    )
    non_empty_blocks = [block for block in raw_blocks if str(block.get("content", "")).strip()]
    readable_chars = sum(len(str(block.get("content", "")).strip()) for block in non_empty_blocks)
    protocol_field_count = sum(len(block.get("metadata", {}).get("protocol_fields", [])) for block in raw_blocks)

    if not raw_blocks:
        issues.append("未提取到任何内容块")
    if readable_chars < MIN_READABLE_CHAR_COUNT:
        issues.append("提取出的有效文本过少，内容可读性不足")
    if total_pages <= 0:
        issues.append("未识别到有效页数或 Sheet 数")

    return {
        "total_units": total_pages,
        "total_blocks": len(raw_blocks),
        "readable_blocks": len(non_empty_blocks),
        "readable_chars": readable_chars,
        "protocol_field_count": protocol_field_count,
    }, issues


def _validate_uploaded_file(uploaded_file, max_size_mb: int) -> dict:
    file_name = (uploaded_file.filename or "").strip()
    ext = os.path.splitext(file_name)[1].lower()
    size_bytes = _file_size_bytes(uploaded_file)
    file_bytes = _read_uploaded_bytes(uploaded_file)
    max_size_bytes = max_size_mb * 1024 * 1024
    checks = {}
    issues = []
    metrics = {
        "total_units": 0,
        "total_blocks": 0,
        "readable_blocks": 0,
        "readable_chars": 0,
        "protocol_field_count": 0,
    }
    detected_type = "unknown"
    temp_file_path = ""

    extension_passed = bool(file_name) and ext in SUPPORTED_EXTS
    checks["extension"] = {
        "passed": extension_passed,
        "actual": ext or "",
        "supported": SUPPORTED_EXTS,
        "message": "扩展名受支持" if extension_passed else f"不支持的文件类型: {ext or '未知'}",
    }
    if not extension_passed:
        issues.append(checks["extension"]["message"])

    size_passed = 0 < size_bytes <= max_size_bytes
    checks["file_size"] = {
        "passed": size_passed,
        "actual_bytes": size_bytes,
        "max_bytes": max_size_bytes,
        "message": "文件大小符合限制" if size_passed else f"文件大小超出限制，当前 {size_bytes} bytes，上限 {max_size_bytes} bytes",
    }
    if not size_passed:
        issues.append(checks["file_size"]["message"])

    authenticity_passed = False
    authenticity_message = "扩展名非法，跳过真实性校验"
    if extension_passed:
        authenticity_passed, detected_type, authenticity_message = _detect_file_authenticity(ext, file_bytes)
        if not authenticity_passed:
            issues.append(authenticity_message)
    checks["authenticity"] = {
        "passed": authenticity_passed,
        "detected_type": detected_type,
        "message": authenticity_message,
    }

    readability_passed = False
    completeness_passed = False
    readability_message = "前置检查未通过，跳过内容可读性分析"
    completeness_message = "前置检查未通过，跳过完整性分析"
    parse_error = None

    if extension_passed and size_passed and authenticity_passed:
        try:
            temp_file_path = _save_validation_temp_file(file_name, file_bytes)
            metrics, content_issues = _analyze_file_content(temp_file_path, file_name)
            readability_passed = metrics["readable_chars"] >= MIN_READABLE_CHAR_COUNT and metrics["readable_blocks"] > 0
            completeness_passed = metrics["total_units"] > 0 and metrics["total_blocks"] > 0
            readability_message = (
                "文件内容可正常读取"
                if readability_passed
                else "可提取内容不足，无法确认文档具备稳定可读性"
            )
            completeness_message = (
                "文件已提取出有效内容块"
                if completeness_passed
                else "未提取出完整内容块或页/Sheet 信息"
            )
            issues.extend(content_issues)
        except Exception as exc:
            parse_error = str(exc)
            issues.append(f"内容解析失败: {parse_error}")
            readability_message = "内容解析失败，无法完成可读性校验"
            completeness_message = "内容解析失败，无法完成完整性校验"
        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                os.remove(temp_file_path)

    checks["readability"] = {
        "passed": readability_passed,
        "message": readability_message,
        "metrics": {
            "readable_blocks": metrics["readable_blocks"],
            "readable_chars": metrics["readable_chars"],
        },
    }
    checks["completeness"] = {
        "passed": completeness_passed,
        "message": completeness_message,
        "metrics": {
            "total_units": metrics["total_units"],
            "total_blocks": metrics["total_blocks"],
            "protocol_field_count": metrics["protocol_field_count"],
        },
    }

    result = {
        "file_name": file_name,
        "extension": ext,
        "size_bytes": size_bytes,
        "max_size_mb": max_size_mb,
        "detected_type": detected_type,
        "valid": all(check["passed"] for check in checks.values()),
        "checks": checks,
        "issues": sorted(set(issue for issue in issues if issue)),
    }
    if parse_error:
        result["parse_error"] = parse_error
    return result


def _collect_uploaded_files() -> list:
    collected = []
    seen = set()
    for uploaded_file in request.files.getlist("files") + [request.files.get("file")]:
        if not uploaded_file or not getattr(uploaded_file, "filename", ""):
            continue
        identity = id(uploaded_file)
        if identity in seen:
            continue
        seen.add(identity)
        collected.append(uploaded_file)
    return collected


@app.route("/api/data/validate_protocol_files", methods=["POST"])
def validate_protocol_files():
    """协议文档基础校验接口。"""
    uploaded_files = _collect_uploaded_files()
    if not uploaded_files:
        return jsonify({
            "code": 400,
            "message": "缺少必要参数: file 或 files"
        }), 400

    max_size_raw = request.form.get("max_size_mb", str(DEFAULT_MAX_FILE_SIZE_MB)).strip()
    try:
        max_size_mb = int(max_size_raw)
        if max_size_mb <= 0:
            raise ValueError
    except ValueError:
        return jsonify({
            "code": 400,
            "message": "max_size_mb 必须为正整数"
        }), 400

    results = [_validate_uploaded_file(uploaded_file, max_size_mb=max_size_mb) for uploaded_file in uploaded_files]
    passed_files = sum(1 for result in results if result["valid"])

    return jsonify({
        "code": 200,
        "message": "success",
        "data": {
            "summary": {
                "total_files": len(results),
                "passed_files": passed_files,
                "failed_files": len(results) - passed_files,
            },
            "results": results,
        }
    })


@app.route("/api/data/upload_split", methods=["POST"])
def upload_split():
    """
    文档上传与智能分割接口

    输入参数 (multipart/form-data):
    - project_id: 项目编号 (必填)
    - file: 上传的文件 (必填)

    响应格式:
    {
        "code": 200,
        "message": "success",
        "data": {
            "task_id": "tsp_001",
            "total_pages": 120,
            "blocks": [...]
        }
    }
    """
    # 获取参数
    project_id = request.form.get("project_id")
    file = request.files.get("file")
    enable_llm_postprocess = request.form.get("enable_llm_postprocess", "true").lower() != "false"

    if not project_id or not file:
        return jsonify({
            "code": 400,
            "message": "缺少必要参数: project_id 或 file"
        }), 400

    # 检查文件类型
    file_name = file.filename
    ext = os.path.splitext(file_name)[1].lower()
    if ext not in SUPPORTED_EXTS:
        return jsonify({
            "code": 400,
            "message": f"不支持的文件类型: {ext}，支持的类型: {', '.join(SUPPORTED_EXTS)}"
        }), 400

    # 保存文件
    file_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4().hex}_{file_name}")
    file.save(file_path)

    try:
        # 处理文件
        raw_blocks, total_pages = process_file(
            file_path,
            project_id,
            file_name,
            enable_llm_postprocess=enable_llm_postprocess,
        )

        # 构建Block对象并保存到数据库
        db_blocks = []
        response_blocks = []

        for idx, block_data in enumerate(raw_blocks):
            # 创建Block对象
            block = Block(
                block_id=0,  # 数据库自动生成
                project_id=project_id,
                file_name=file_name,
                page_num=block_data["page_num"],
                content=block_data["content"],
                block_type=block_data["type"],
                metadata=block_data.get("metadata", {})
            )
            db_blocks.append(block)

            # 构建响应数据
            response_blocks.append({
                "block_id": idx + 1,
                "page_num": block_data["page_num"],
                "content": block_data["content"],
                "type": block_data["type"],
                "protocol_fields": block_data.get("metadata", {}).get("protocol_fields", []),
            })

        # 保存到数据库
        db_block_ids = []
        try:
            db_block_ids = mysql_client.insert_blocks(db_blocks)
        except Exception as db_error:
            # 数据库保存失败不影响响应，记录日志
            print(f"数据库保存失败: {str(db_error)}")

        # 更新响应中的block_id为数据库ID
        if db_block_ids:
            for i, db_id in enumerate(db_block_ids):
                if i < len(response_blocks):
                    response_blocks[i]["block_id"] = db_id

        # 保存中间结果到文件存储
        task_id = f"tsp_{int(time.time())}_{uuid.uuid4().hex[:8]}"
        try:
            blocks_for_store = [
                {
                    "block_id": response_blocks[i]["block_id"],
                    "page_num": b["page_num"],
                    "content": b["content"],
                    "type": b["type"],
                    "metadata": raw_blocks[i].get("metadata", {})
                }
                for i, b in enumerate(response_blocks)
            ]
            file_store.save_blocks(project_id, blocks_for_store, file_name=f"{task_id}_blocks.json")
        except Exception as store_error:
            print(f"文件存储失败: {str(store_error)}")

        # 统计表格和文本数量
        table_count = sum(1 for b in response_blocks if b["type"] == "table")
        text_count = len(response_blocks) - table_count
        protocol_field_count = sum(len(b.get("protocol_fields", [])) for b in response_blocks)

        return jsonify({
            "code": 200,
            "message": "success",
            "data": {
                "task_id": task_id,
                "total_pages": total_pages,
                "total_blocks": len(response_blocks),
                "table_count": table_count,
                "text_count": text_count,
                "protocol_field_count": protocol_field_count,
                "blocks": response_blocks
            }
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            "code": 500,
            "message": f"处理失败: {str(e)}"
        }), 500

    finally:
        # 清理临时文件
        if os.path.exists(file_path):
            os.remove(file_path)


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "healthy"})


if __name__ == "__main__":
    # 初始化数据库表
    try:
        mysql_client.init_tables()
    except Exception as e:
        print(f"数据库初始化失败: {str(e)}")

    app.run(host="0.0.0.0", port=5001, debug=True)
