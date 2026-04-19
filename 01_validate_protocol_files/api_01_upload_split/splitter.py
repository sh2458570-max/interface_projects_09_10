# splitter.py赵清欢的“读文件+分块”模块，独立成一个文件
import os
from typing import List, Optional
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownTextSplitter,
    PythonCodeTextSplitter,
    CharacterTextSplitter,
)

from langchain_core.documents import Document

import pdfplumber


class DocumentSplitter:
    def __init__(self):
        self.splitters = self._initialize_splitters()

    def _initialize_splitters(self):
        return {
            "py": PythonCodeTextSplitter(chunk_size=400, chunk_overlap=50),
            "python": PythonCodeTextSplitter(chunk_size=400, chunk_overlap=50),
            "java": CharacterTextSplitter(chunk_size=400, chunk_overlap=50, separator="\n\n"),
            "cpp": CharacterTextSplitter(chunk_size=400, chunk_overlap=50, separator="\n\n"),
            "js": CharacterTextSplitter(chunk_size=400, chunk_overlap=50, separator="\n\n"),
            "md": MarkdownTextSplitter(chunk_size=800, chunk_overlap=100),
            "markdown": MarkdownTextSplitter(chunk_size=800, chunk_overlap=100),
            "json": CharacterTextSplitter(chunk_size=500, chunk_overlap=50, separator="\n"),
            "yaml": CharacterTextSplitter(chunk_size=500, chunk_overlap=50, separator="\n"),
            "yml": CharacterTextSplitter(chunk_size=500, chunk_overlap=50, separator="\n"),
            "xml": CharacterTextSplitter(chunk_size=500, chunk_overlap=50, separator="\n"),
            "pdf": RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=100,
                separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", ";", "；", " ", ""],
            ),
            "txt": RecursiveCharacterTextSplitter(
                chunk_size=800,
                chunk_overlap=80,
                separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
            ),
            "docx": RecursiveCharacterTextSplitter(
                chunk_size=800,
                chunk_overlap=80,
                separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
            ),
            "general": RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=50,
                separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
            ),
        }

    def get_splitter(self, doc_type: Optional[str] = None, file_path: Optional[str] = None):
        if doc_type and doc_type in self.splitters:
            return self.splitters[doc_type]
        if file_path:
            ext = os.path.splitext(file_path)[1].lower().lstrip(".")
            if ext in self.splitters:
                return self.splitters[ext]
        return self.splitters["general"]

    def split_to_documents(self, content: str, *, file_path: Optional[str] = None, doc_type: Optional[str] = None) -> List[Document]:
        splitter = self.get_splitter(doc_type=doc_type, file_path=file_path)

        if hasattr(splitter, "split_text"):
            chunks = splitter.split_text(content)
        else:
            docs = splitter.create_documents([content])
            chunks = [d.page_content for d in docs]

        file_name = os.path.basename(file_path) if file_path else "unknown"
        file_type = os.path.splitext(file_name)[1].lstrip(".") if file_path else "unknown"

        documents = []
        total = len(chunks)
        for idx, chunk in enumerate(chunks):
            if not chunk.strip():
                continue
            documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "file_name": file_name,
                        "file_type": file_type,
                        "chunk_id": idx,
                        "total_chunks": total,
                    },
                )
            )
        return documents


def load_file_content(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext in [".txt", ".md", ".py", ".java", ".js", ".cpp", ".json", ".xml", ".yaml", ".yml"]:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif ext == ".pdf":
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        return text
    elif ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")
